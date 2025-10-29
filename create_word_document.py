#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_spiritual_homework():
    # إنشاء مستند جديد
    doc = Document()
    
    # إعداد الخط العربي
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # العنوان الرئيسي
    title = doc.add_heading('📖 الواجب الروحي الأسبوعي', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('✝️ تأملات روحية من رسالة تيموثاوس الأولى - الإصحاح الخامس', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # إضافة خط فاصل
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # قسم القراءة المطلوبة
    doc.add_heading('🎯 القراءة المطلوبة:', level=2)
    reading_list = doc.add_paragraph()
    reading_list.add_run('• ').bold = True
    reading_list.add_run('رسالة تيموثاوس الأولى - الإصحاح 5').bold = True
    reading_list.add_run('\n• ').bold = True
    reading_list.add_run('كتاب "الرب يسوع خادم النفوس" - الباب الأول (ص 13-28)').bold = True
    
    # إضافة خط فاصل
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # تأملات من رسالة تيموثاوس
    doc.add_heading('📝 تأملاتي من رسالة تيموثاوس الأولى - الإصحاح 5:', level=2)
    
    # الآيات المؤثرة
    doc.add_heading('💭 الآيات المؤثرة:', level=3)
    verse_para = doc.add_paragraph()
    verse_para.add_run('"لا توبخ شيخاً، بل عظه كأب، والشبان كإخوة، والعجائز كأمهات، والفتيات كأخوات بكل طهارة"').italic = True
    verse_para.add_run('\n(تيموثاوس الأولى 5: 1-2)')
    
    # التأملات الروحية
    doc.add_heading('🌟 التأملات الروحية:', level=3)
    
    # تأمل 1
    doc.add_heading('1. احترام الكبار والشيوخ:', level=4)
    reflection1 = doc.add_paragraph()
    reflection1.add_run('• الوصية واضحة: لا نوبخ الشيوخ، بل نعظهم كآباء\n')
    reflection1.add_run('• هذا يعلمنا الاحترام والتقدير للكبار في السن\n')
    reflection1.add_run('• حتى لو كانوا مخطئين، نتعامل معهم بحكمة ومحبة')
    
    # تأمل 2
    doc.add_heading('2. العلاقات العائلية الروحية:', level=4)
    reflection2 = doc.add_paragraph()
    reflection2.add_run('• نتعامل مع الشبان كإخوة - علاقة مساواة ومحبة\n')
    reflection2.add_run('• العجائز كأمهات - رعاية وحنان\n')
    reflection2.add_run('• الفتيات كأخوات بكل طهارة - حماية وطهارة')
    
    # تأمل 3
    doc.add_heading('3. رعاية الأرامل:', level=4)
    reflection3 = doc.add_paragraph()
    reflection3.add_run('• اهتمام خاص بالأرامل الحقيقيات\n')
    reflection3.add_run('• مساعدتهم مادياً وروحياً\n')
    reflection3.add_run('• عدم إهمال من هم في حاجة')
    
    # إضافة خط فاصل
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # تأملات من الكتاب
    doc.add_heading('📚 تأملاتي من كتاب "الرب يسوع خادم النفوس" - الباب الأول:', level=2)
    
    doc.add_heading('🎯 الأفكار الروحية المستفادة:', level=3)
    
    # فكرة 1
    doc.add_heading('1. يسوع الخادم الحقيقي:', level=4)
    idea1 = doc.add_paragraph()
    idea1.add_run('• المسيح لم يأتِ ليُخدم بل ليخدم\n')
    idea1.add_run('• خدمة الآخرين هي طريق القداسة\n')
    idea1.add_run('• كل عمل صغير نعمله للآخرين هو خدمة للمسيح')
    
    # فكرة 2
    doc.add_heading('2. محبة الآخرين:', level=4)
    idea2 = doc.add_paragraph()
    idea2.add_run('• المحبة الحقيقية تظهر في الأفعال وليس الكلام فقط\n')
    idea2.add_run('• خدمة الآخرين بدون انتظار مقابل\n')
    idea2.add_run('• وضع احتياجات الآخرين قبل احتياجاتنا')
    
    # فكرة 3
    doc.add_heading('3. التواضع في الخدمة:', level=4)
    idea3 = doc.add_paragraph()
    idea3.add_run('• الخدمة الحقيقية تحتاج تواضع\n')
    idea3.add_run('• عدم التباهي بما نعمله للآخرين\n')
    idea3.add_run('• خدمة في الخفاء كما في العلن')
    
    # إضافة خط فاصل
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # التطبيق العملي
    doc.add_heading('🎯 التطبيق العملي - ما سأجاهد في تنفيذه:', level=2)
    
    doc.add_heading('✅ هذا الأسبوع سأعمل على:', level=3)
    
    # تطبيق 1
    doc.add_heading('1. احترام الكبار:', level=4)
    practical1 = doc.add_paragraph()
    practical1.add_run('• سأتحدث مع الشيوخ باحترام وتقدير\n')
    practical1.add_run('• سأستمع لنصائحهم بحب\n')
    practical1.add_run('• سأساعدهم في احتياجاتهم')
    
    # تطبيق 2
    doc.add_heading('2. خدمة الآخرين:', level=4)
    practical2 = doc.add_paragraph()
    practical2.add_run('• سأساعد شخص واحد على الأقل يومياً\n')
    practical2.add_run('• سأخدم في الكنيسة بفرح\n')
    practical2.add_run('• سأضع احتياجات الآخرين قبل راحتي')
    
    # تطبيق 3
    doc.add_heading('3. التواضع:', level=4)
    practical3 = doc.add_paragraph()
    practical3.add_run('• سأخدم في الخفاء بدون تباهي\n')
    practical3.add_run('• سأشكر الله على كل فرصة للخدمة\n')
    practical3.add_run('• سأطلب من الله أن يعطيني قلب خادم')
    
    # إضافة خط فاصل
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # صلاة الختام
    doc.add_heading('🙏 صلاة الختام:', level=2)
    prayer_para = doc.add_paragraph()
    prayer_para.add_run('"يا رب يسوع، أعطني قلب خادم مثلك. علمني أن أحترم الكبار وأخدم الجميع بمحبة وتواضع. ساعدني أن أطبق كلمتك في حياتي اليومية. آمين."').italic = True
    prayer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # إضافة خط فاصل
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # التوقيع
    doc.add_paragraph('📅 تاريخ الإنجاز: ___________')
    doc.add_paragraph('✍️ التوقيع: ___________')
    
    # آية ختامية
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    final_verse = doc.add_paragraph()
    final_verse.add_run('"لأن ابن الإنسان لم يأتِ ليُخدم بل ليخدم وليبذل نفسه فدية عن كثيرين"').italic = True
    final_verse.add_run('\n(مرقس 10: 45)')
    final_verse.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # حفظ الملف
    doc.save('الواجب_الروحي_الأسبوعي.docx')
    print("تم إنشاء ملف Word بنجاح: الواجب_الروحي_الأسبوعي.docx")

if __name__ == "__main__":
    try:
        create_spiritual_homework()
    except ImportError:
        print("تحتاج إلى تثبيت مكتبة python-docx:")
        print("pip install python-docx")
    except Exception as e:
        print(f"حدث خطأ: {e}")
